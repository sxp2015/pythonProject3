import docx

doc = docx.Document('input123.docx')

doc_length = len(doc.paragraphs)

print('文档中有', doc_length, '个段落')

for p in doc.paragraphs:
    # print('Paragraph', p.text)
    for i in p.runs:
        print('Paragraph Run:', i.text)

# table = doc.add_table(rows=1, cols=2)
# row = table.rows[0]
# row.cells[0].text = '第一列'
# row.cells[1].text = '第二列'
#
# doc.add_paragraph('表格内容1', style='List Bullet')
# doc.add_paragraph('表格内容2', style='List Bullet')
# doc.add_paragraph('表格内容3', style='List Bullet')
#
# table.style = ''
# doc.save('output.docx')