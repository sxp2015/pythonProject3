# 导入 docx 模块
import docx

# 读取指定路径下的 Word 文档
doc = docx.Document(r'C:\Users\admin\PycharmProjects\pythonProject3\study\docx\test\01.docx')

# 获取文档中段落的数量
paragraphs_count = len(doc.paragraphs)

# 遍历每一个段落，输出它的内容
for i, paragraph in enumerate(doc.paragraphs):
    print(f'第{i + 1}段的内容是：{paragraph.text}')

# 获取某个特定段落中的特定运行（即文本格式）的文本内容，并将其保存在变量 run1 和 run2 中
run1 = doc.paragraphs[1].runs[2].text
run2 = doc.paragraphs[1].runs[4].text

# 修改样式
style_change_before = doc.paragraphs[1].style
style_change_before_text = doc.paragraphs[1].text
print('style_change_before： ', style_change_before)
print('style_change_before_text： ', style_change_before_text)
doc.paragraphs[1].runs[2].style = 'Char'

doc.save('style_change.docx')


# 打印文档的段落数、run1 的文本内容和 run2 的文本内容
# print('len', len(doc.paragraphs))
# print('run1是:', run1)
# print('run2是:', run2)
