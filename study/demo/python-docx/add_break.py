from time import sleep

from office import word
from docx import Document, shared
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_BREAK

# 图片路径
img_path = r'C:\Users\admin\PycharmProjects\pythonProject3\study\img\avatar.png'
img_path2 = r'E:\Python\2022\pythonProject3\study\img\avatar.png'
# 创建 Document 对象，设置页面大小和方向
doc = Document()
sectionItem = doc.sections[0]
sectionItem.page_width = shared.Cm(21)  # 页面宽度设置为A4纵向方向的宽度
sectionItem.page_height = shared.Cm(29.7)  # 页面高度设置为A4纵向方向的高度
sectionItem.orientation = WD_ORIENTATION.PORTRAIT  # 纵向方向

# 添加标题和段落
doc.add_heading('测试添加分页符', level=1)
doc.add_paragraph('这是第一页的内容')

# 添加分页符
doc.paragraphs[0].runs[0].add_break(WD_BREAK.PAGE)

# 添加标题和段落
doc.add_paragraph('这是第二页的内容', style='Title')

doc.add_picture(img_path2, height=shared.Cm(4), width=shared.Inches(1.5))

# 保存文档
doc.save('twoPageA4.docx')

sleep(2)

# word转pdf
word.docx2pdf(path='twoPageA4.docx')

