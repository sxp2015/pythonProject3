import docx
# 创建文档对象
doc = docx.Document()

# 添加标题
doc.add_heading('请在此输入标题0', 0)
doc.add_heading('请在此输入标题1', 1)
doc.add_heading('请在此输入标题2', 2)
doc.add_heading('请在此输入标题3', 3)
doc.add_heading('请在此输入标题4', 4)

# 添加段落
doc.add_paragraph('hello world', 'Title')
paraObj2 = doc.add_paragraph('这是添加的第二段内容')
paraObj3 = doc.add_paragraph('其他的内容在这里')

# 追加文本
paraObj2.add_run('这是追加到，第二段内容后面的文字。。。。')

# 保存
doc.save('multipleParagraphs.docx')

""""
add_paragraph()和add_run()都接收可选的第二个参数，它是表示Paragraph
或Run对象样式的字符串。例如：
>>> doc.add_paragraph('Hello, world!', 'Title')
这一行添加了一段，文本是Hello, world!，样式是Title
"""
